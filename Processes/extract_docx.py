#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 文档内容提取工具
用于提取 .docx 文件的文本内容
"""

import sys
from pathlib import Path

def extract_docx(docx_path, output_path=None):
    """
    提取 Word 文档内容到文本文件
    
    Args:
        docx_path: Word 文档路径
        output_path: 输出文本文件路径（可选）
    """
    try:
        from docx import Document
    except ImportError:
        print("错误: 需要安装 python-docx 库")
        print("请运行: pip3 install python-docx -i https://pypi.tuna.tsinghua.edu.cn/simple")
        return False
    
    try:
        # 读取 Word 文档
        doc = Document(docx_path)
        
        # 提取所有段落文本
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():  # 跳过空段落
                full_text.append(para.text)
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    full_text.append(row_text)
        
        content = '\n\n'.join(full_text)
        
        # 输出到文件或打印
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            print(f"✅ 内容已提取到: {output_path}")
        else:
            print(content)
        
        return True
        
    except Exception as e:
        print(f"❌ 提取失败: {e}")
        return False

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python3 extract_docx.py <docx文件路径> [输出文件路径]")
        print("\n示例:")
        print("  python3 extract_docx.py document.docx")
        print("  python3 extract_docx.py document.docx output.txt")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not Path(docx_path).exists():
        print(f"❌ 文件不存在: {docx_path}")
        sys.exit(1)
    
    extract_docx(docx_path, output_path)
